# -*- coding: utf-8 -*-
"""
양식 폴더에서 사업계획서 양식 데이터를 추출하여 JSON으로 저장
- .docx: python-docx 사용
- .doc: pywin32 Word COM 사용
"""
import json
import os
import sys

# 콘솔 인코딩 설정
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx

FORMS_DIR = r"C:\Users\chado\OneDrive\문서\양식 폴더"
OUTPUT_DIR = r"C:\project\Draft\docs"

# 각 양식별 핵심 사업계획서 파일
FORM_FILES = {
    "yebi-chogi": {
        "name": "예비창업패키지",
        "file": r"01_예비창업패키지\[별첨 1] 2025년도 예비창업패키지 사업계획서 양식.docx"
    },
    "chogi": {
        "name": "초기창업패키지",
        "file": r"02_초기창업패키지\[별첨 1] 2025년도 초기창업패키지 사업계획서 양식.doc"
    },
    "student-300": {
        "name": "학생창업유망팀300",
        "file": r"03_학생창업유망팀300\양식1. 사업계획서_2025 학생 창업유망팀 300+.doc"
    },
    "saengae-chungnyeon": {
        "name": "생애최초청년창업",
        "file": r"04_생애최초청년창업\vertopal.com_[별첨1] 생애최초 청년창업 지원사업 사업계획서 양식.doc"
    },
    "oneul-jeongtong": {
        "name": "오늘전통",
        "file": r"05_오늘전통\vertopal.com_붙임1. 2025 오늘전통 청년 예비창업 공모전 제출서류 양식.doc"
    },
    "gyeonggi-g-star": {
        "name": "경기G스타오디션",
        "file": r"06_경기G스타오디션\vertopal.com_붙임2.+참가신청+양식(2026+경기+창업+공모(G스타+오디션-예비초기리그)).doc"
    }
}

def extract_doc_with_win32(doc_path):
    """Word COM을 사용해서 .doc/.docx 파일 읽기"""
    import win32com.client
    import pythoncom

    pythoncom.CoInitialize()

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(doc_path)

        result = {
            "paragraphs": [],
            "tables": [],
            "full_text": ""
        }

        # 전체 텍스트
        result["full_text"] = doc.Content.Text

        # 단락 추출
        for para in doc.Paragraphs:
            text = para.Range.Text.strip()
            if text and text != '\r':
                result["paragraphs"].append(text)

        # 테이블 추출
        for i, table in enumerate(doc.Tables):
            table_data = []
            try:
                for row in table.Rows:
                    row_data = []
                    for cell in row.Cells:
                        cell_text = cell.Range.Text.strip().replace('\r\x07', '').replace('\r', ' ')
                        row_data.append(cell_text)
                    if any(row_data):
                        table_data.append(row_data)
            except:
                pass
            if table_data:
                result["tables"].append({
                    "index": i,
                    "rows": table_data
                })

        doc.Close(False)
        word.Quit()

        return result

    except Exception as e:
        try:
            word.Quit()
        except:
            pass
        return {"error": str(e)}
    finally:
        pythoncom.CoUninitialize()

def extract_docx_content(doc_path):
    """python-docx로 .docx 파일 읽기"""
    try:
        doc = docx.Document(doc_path)

        result = {
            "paragraphs": [],
            "tables": [],
            "full_text": ""
        }

        # 단락 추출
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                result["paragraphs"].append(text)

        result["full_text"] = "\n".join(result["paragraphs"])

        # 테이블 추출
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                if any(row_data):
                    table_data.append(row_data)
            if table_data:
                result["tables"].append({
                    "index": i,
                    "rows": table_data
                })

        return result

    except Exception as e:
        return {"error": str(e)}

def extract_doc_content(doc_path):
    """파일 확장자에 따라 적절한 방법으로 추출"""
    if doc_path.lower().endswith('.docx'):
        return extract_docx_content(doc_path)
    else:
        return extract_doc_with_win32(doc_path)

def analyze_form_structure(content):
    """양식 구조 분석 - 섹션, 필드, 배점 등"""
    structure = {
        "evaluation_items": [],
        "required_fields": [],
        "psst_sections": {},
        "key_sections": []
    }

    full_text = content.get("full_text", "")

    # PSST 키워드 검색
    psst_markers = {
        "problem": ["문제 인식", "Problem", "창업아이템의 필요성", "시장현황"],
        "solution": ["실현 가능성", "Solution", "개발 계획", "차별성", "경쟁력"],
        "scaleup": ["성장전략", "Scale-up", "사업화 전략", "시장 규모", "비즈니스 모델"],
        "team": ["팀 구성", "Team", "대표자", "조직 역량", "인력 구성"]
    }

    for section, keywords in psst_markers.items():
        for keyword in keywords:
            if keyword in full_text:
                if section not in structure["psst_sections"]:
                    structure["psst_sections"][section] = []
                structure["psst_sections"][section].append(keyword)

    # 테이블에서 평가항목/배점 추출
    for table in content.get("tables", []):
        for row in table.get("rows", []):
            row_text = " ".join(str(cell) for cell in row)

            # 배점 정보 찾기
            if '점' in row_text or '배점' in row_text:
                structure["evaluation_items"].append(row)

    # 주요 섹션 텍스트 추출
    for para in content.get("paragraphs", []):
        if any(k in para for keywords in psst_markers.values() for k in keywords):
            if len(para) < 200:  # 제목성 텍스트만
                structure["key_sections"].append(para)

    return structure

def main():
    print("=" * 60)
    print("양식 데이터 추출 시작")
    print("=" * 60)

    all_forms_data = {}

    for form_id, form_info in FORM_FILES.items():
        file_path = os.path.join(FORMS_DIR, form_info["file"])
        print(f"\n[{form_info['name']}] 처리 중...")

        if not os.path.exists(file_path):
            print(f"  - 파일 없음: {file_path}")
            continue

        content = extract_doc_content(file_path)

        if "error" in content:
            print(f"  - 오류: {content['error']}")
            continue

        structure = analyze_form_structure(content)

        all_forms_data[form_id] = {
            "name": form_info["name"],
            "file": form_info["file"],
            "paragraph_count": len(content.get("paragraphs", [])),
            "table_count": len(content.get("tables", [])),
            "structure": structure,
            "sample_paragraphs": content.get("paragraphs", [])[:20],  # 처음 20개 단락
            "tables": content.get("tables", [])[:5]  # 처음 5개 테이블
        }

        print(f"  - 단락: {len(content.get('paragraphs', []))}개")
        print(f"  - 테이블: {len(content.get('tables', []))}개")
        print(f"  - PSST 섹션: {list(structure['psst_sections'].keys())}")

    # JSON으로 저장
    output_path = os.path.join(OUTPUT_DIR, "extracted_form_data.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_forms_data, f, ensure_ascii=False, indent=2)

    print(f"\n\n저장 완료: {output_path}")
    print("=" * 60)

if __name__ == "__main__":
    main()
